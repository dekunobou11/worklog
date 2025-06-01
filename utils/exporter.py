# utils/exporter.py
import os
import zipfile
import io
from flask import send_file
from docx import Document
from docx.shared import Inches
import markdown2

def export_as_zip(memos, format):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zipf:
        if format == 'word':
            doc = Document()
            for memo in memos:
                doc.add_heading(memo['created_at'], level=1)
                doc.add_paragraph(memo['content'])
                doc.add_paragraph("タグ: " + memo['tag'])
                if memo['image']:
                    image_path = os.path.join('uploads', memo['image'])
                    if os.path.exists(image_path):
                        try:
                            doc.add_picture(image_path, width=Inches(4))
                        except Exception as e:
                            doc.add_paragraph(f"[画像読み込みエラー: {memo['image']}]")
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            zipf.writestr("memos.docx", doc_bytes.read())

        else:  # markdown
            combined_md = ""
            for memo in memos:
                combined_md += f"# {memo['created_at']}\n\n"
                combined_md += memo['content'] + "\n\n"
                combined_md += f"**タグ**: {memo['tag']}\n\n"
                if memo['image']:
                    combined_md += f"![image](images/{memo['image']})\n\n"
            zipf.writestr("memos.md", combined_md.encode('utf-8'))

        # 画像ファイルを追加
        for memo in memos:
            if memo['image']:
                image_path = os.path.join('uploads', memo['image'])
                if os.path.exists(image_path):
                    zipf.write(image_path, arcname=os.path.join('images', memo['image']))

    zip_buffer.seek(0)
    return send_file(zip_buffer, mimetype='application/zip', as_attachment=True, download_name='memos_export.zip')
